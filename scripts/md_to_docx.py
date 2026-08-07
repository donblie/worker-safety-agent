"""
将参赛说明文档.md 转换为 Word 文档 (.docx)
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing

def add_styled_paragraph(doc, text, style=None, bold=False, font_size=None, color=None,
                         alignment=None, font_name=None, before=0, after=0):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p, before=before, after=after)
    return p

def add_code_block(doc, text, font_size=8):
    """添加代码块(ASCII图)"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

def add_bullet(doc, text, level=0, bold_prefix=None):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10.5)
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(10.5)
    else:
        # 处理markdown格式
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 去掉加粗标记（后面再单独处理）
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_table_from_data(doc, headers, rows, col_widths=None):
    """从数据创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        # 清理markdown加粗标记
        header_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', header)
        run = p.add_run(header_clean)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 表头背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1B5E9B')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            # 处理加粗: **text** → 拆分为多个run
            parts = re.split(r'(\*\*.+?\*\*)', str(cell_text))
            for part in parts:
                bold_match = re.match(r'\*\*(.+?)\*\*', part)
                if bold_match:
                    run = p.add_run(bold_match.group(1))
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table

def convert_md_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过空行
        if not line:
            i += 1
            continue

        # ── 标题 ──
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_styled_paragraph(doc, text, bold=True, font_size=22,
                                font_name='微软雅黑', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                before=12, after=6)
        elif line.startswith('## '):
            text = line[3:].strip()
            add_styled_paragraph(doc, text, bold=True, font_size=16,
                                font_name='微软雅黑', before=12, after=6,
                                color=(0x1B, 0x5E, 0x9B))
        elif line.startswith('### '):
            text = line[4:].strip()
            add_styled_paragraph(doc, text, bold=True, font_size=13,
                                font_name='微软雅黑', before=8, after=4,
                                color=(0x2C, 0x3E, 0x50))

        # ── 引用块 ──
        elif line.startswith('> '):
            text = line[2:].strip()
            # 去掉markdown加粗标记
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_styled_paragraph(doc, text, font_size=10, color=(0x66, 0x66, 0x66),
                                font_name='楷体', before=4, after=4)

        # ── 分隔线 ──
        elif line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── 表格（检测 | 开头的行）──
        elif line.startswith('|') and line.strip().endswith('|'):
            # 收集连续的表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # 回退，外层会+1

            # 过滤分隔行（如 |------|------|）
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if all(re.match(r'^[-:]+$', c) for c in cells):
                    continue  # 跳过分隔行
                data_rows.append(cells)

            if len(data_rows) >= 2:
                headers = data_rows[0]
                rows = data_rows[1:]
                add_table_from_data(doc, headers, rows)

        # ── 代码块 ──
        elif line.strip() == '```':
            code_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                code_lines.append(lines[i].rstrip())
                i += 1
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))

        # ── 无序列表 ──
        elif line.startswith('- ') or line.startswith('  - '):
            indent = len(line) - len(line.lstrip())
            text = line.lstrip('- ').strip()
            # 处理加粗
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if indent > 0:
                p.paragraph_format.left_indent = Cm(1.27 * (indent // 2))

        # ── 普通段落 ──
        else:
            # 清理markdown格式
            text = line
            # 粗体 **text**
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # 行内代码 `text`
            text = re.sub(r'`(.+?)`', r'\1', text)

            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            set_paragraph_spacing(p, before=3, after=3)

        i += 1

    # ── 保存 ──
    doc.save(docx_path)
    print(f"✅ Word文档已生成: {docx_path}")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base_dir, 'docs', '参赛说明文档.md')
    docx_path = os.path.join(base_dir, 'docs', '参赛说明文档.docx')
    convert_md_to_docx(md_path, docx_path)
